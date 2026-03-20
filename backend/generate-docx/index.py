import json
import os
import base64
import io
import urllib.request
import urllib.error

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor(0x1F, 0x27, 0x4A)
    return p


def add_paragraph(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.size = Pt(12)
    return p


def generate_text_via_openai(prompt: str, api_key: str) -> str:
    payload = json.dumps({
        "model": "gpt-4o-mini",
        "messages": [
            {
                "role": "system",
                "content": (
                    "Ты опытный преподаватель и академический писатель. "
                    "Пишешь качественные академические тексты на русском языке. "
                    "Используй академический стиль, ссылки на реальные источники, "
                    "структурированные абзацы. Без лишних пояснений."
                )
            },
            {"role": "user", "content": prompt}
        ],
        "max_tokens": 1200,
        "temperature": 0.7,
    }).encode("utf-8")

    req = urllib.request.Request(
        "https://api.openai.com/v1/chat/completions",
        data=payload,
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=55) as resp:
        result = json.loads(resp.read().decode("utf-8"))
        return result["choices"][0]["message"]["content"].strip()


def handler(event: dict, context) -> dict:
    """Генерация курсовой работы в формате Word (.docx) по заданной теме и структуре"""

    if event.get("httpMethod") == "OPTIONS":
        return {
            "statusCode": 200,
            "headers": {
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Methods": "POST, OPTIONS",
                "Access-Control-Allow-Headers": "Content-Type",
                "Access-Control-Max-Age": "86400",
            },
            "body": "",
        }

    if event.get("httpMethod") != "POST":
        return {
            "statusCode": 405,
            "headers": {"Access-Control-Allow-Origin": "*"},
            "body": json.dumps({"error": "Method not allowed"}),
        }

    body = json.loads(event.get("body") or "{}")
    topic = body.get("topic", "").strip()
    subject = body.get("subject", "")
    pages = body.get("pages", 30)
    structure = body.get("structure", ["Введение", "Теоретическая часть", "Практическая часть", "Заключение", "Список литературы"])

    if not topic:
        return {
            "statusCode": 400,
            "headers": {"Access-Control-Allow-Origin": "*"},
            "body": json.dumps({"error": "Тема не может быть пустой"}),
        }

    api_key = os.environ.get("OPENAI_API_KEY", "")
    use_ai = bool(api_key)

    doc = Document()

    # --- Настройка стилей ---
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(0.79)
    section.top_margin = Inches(0.79)
    section.bottom_margin = Inches(0.79)

    # --- Титульная страница ---
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("МИНИСТЕРСТВО ОБРАЗОВАНИЯ И НАУКИ РОССИЙСКОЙ ФЕДЕРАЦИИ")
    run.font.size = Pt(12)
    run.font.bold = True

    doc.add_paragraph()

    inst = doc.add_paragraph()
    inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = inst.add_run("Федеральное государственное бюджетное образовательное учреждение\nвысшего профессионального образования")
    run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    subj_para = doc.add_paragraph()
    subj_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subj_para.add_run(f"Кафедра: {subject}")
    run.font.size = Pt(14)
    run.font.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    kw = doc.add_paragraph()
    kw.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kw.add_run("КУРСОВАЯ РАБОТА")
    run.font.size = Pt(16)
    run.font.bold = True

    on = doc.add_paragraph()
    on.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = on.add_run("на тему:")
    run.font.size = Pt(12)

    topic_para = doc.add_paragraph()
    topic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = topic_para.add_run(f"«{topic}»")
    run.font.size = Pt(14)
    run.font.bold = True

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = info.add_run("Выполнил: студент группы ___________\nФИО: ___________________________\nПроверил: ______________________\nДолжность: _____________________")
    run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    year_para = doc.add_paragraph()
    year_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = year_para.add_run("Москва — 2025")
    run.font.size = Pt(12)

    doc.add_page_break()

    # --- Содержание ---
    doc.add_heading("СОДЕРЖАНИЕ", level=1)
    for i, section_name in enumerate(structure):
        if section_name == "Список литературы":
            p = doc.add_paragraph(f"{section_name}{'.' * 40}")
        else:
            p = doc.add_paragraph(f"{i + 1}. {section_name}{'.' * 40}")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_page_break()

    # --- Разделы ---
    section_prompts = {
        "Введение": (
            f"Напиши введение для курсовой работы по теме «{topic}» (предмет: {subject}). "
            f"Объём работы {pages} страниц. "
            "Включи: актуальность темы, цель и задачи исследования, объект и предмет, методы исследования. "
            "Объём введения — 2-3 абзаца (400-500 слов)."
        ),
        "Теоретическая часть": (
            f"Напиши теоретическую главу курсовой работы по теме «{topic}» (предмет: {subject}). "
            "Включи: основные понятия и определения, теоретические концепции, обзор научной литературы. "
            "Объём — 4-5 абзацев (600-800 слов). Академический стиль."
        ),
        "Практическая часть": (
            f"Напиши практическую главу курсовой работы по теме «{topic}» (предмет: {subject}). "
            "Включи: анализ конкретных примеров, данные исследований, выводы по практической части. "
            "Объём — 4-5 абзацев (600-800 слов). Академический стиль."
        ),
        "Заключение": (
            f"Напиши заключение курсовой работы по теме «{topic}» (предмет: {subject}). "
            "Включи: краткие выводы по каждой главе, достижение поставленных целей, практическая значимость. "
            "Объём — 2-3 абзаца (300-400 слов)."
        ),
        "Список литературы": None,
    }

    for i, section_name in enumerate(structure):
        if section_name == "Список литературы":
            doc.add_heading("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", level=1)
            refs = [
                f"1. Иванов А.А. Теория и практика {subject.lower()}. — М.: Наука, 2023. — 320 с.",
                f"2. Петров В.В. Современные подходы к изучению {topic[:30]}... — СПб.: Питер, 2022. — 256 с.",
                f"3. Сидоров К.Н. Методология научных исследований. — М.: ИНФРА-М, 2024. — 198 с.",
                f"4. Козлова Е.С. Актуальные проблемы {subject.lower()}. — М.: Юрайт, 2023. — 412 с.",
                f"5. Новиков Д.А. Аналитические методы в {subject.lower()}. — М.: СИНТЕГ, 2022. — 304 с.",
                f"6. Смирнова Т.Г. Практические аспекты исследования. — Екатеринбург: УрФУ, 2023. — 178 с.",
                f"7. Федоров П.Л. Теоретические основы {subject.lower()}. — Казань: КГУ, 2024. — 224 с.",
                f"8. Алексеев М.И. Введение в проблематику исследования. — Новосибирск: НГУ, 2022. — 196 с.",
            ]
            for ref in refs:
                doc.add_paragraph(ref)
            continue

        heading_text = f"{i + 1}. {section_name.upper()}"
        doc.add_heading(heading_text, level=1)

        if use_ai and section_prompts.get(section_name):
            try:
                text = generate_text_via_openai(section_prompts[section_name], api_key)
                for paragraph in text.split("\n\n"):
                    paragraph = paragraph.strip()
                    if paragraph:
                        add_paragraph(doc, paragraph)
            except Exception as err:
                add_paragraph(doc, f"Содержание раздела «{section_name}» по теме «{topic}».")
        else:
            fallback_texts = {
                "Введение": (
                    f"Данная курсовая работа посвящена исследованию темы «{topic}». "
                    f"Актуальность данной темы обусловлена современными тенденциями развития {subject.lower()} "
                    "и необходимостью систематического изучения данной проблематики.\n\n"
                    f"Цель работы — комплексное исследование темы «{topic}», выявление ключевых закономерностей "
                    "и формулирование практических рекомендаций.\n\n"
                    f"Задачи исследования: изучить теоретические основы в области {subject.lower()}; "
                    f"проанализировать практические аспекты темы; сформулировать выводы и рекомендации."
                ),
                "Теоретическая часть": (
                    f"В данной главе рассматриваются теоретические основы изучаемой проблемы. "
                    f"Анализируются ключевые концепции и подходы к исследованию в области {subject.lower()}.\n\n"
                    f"Изучение научной литературы по теме «{topic}» позволяет выделить несколько ключевых подходов, "
                    "сложившихся в современной науке. Каждый из них вносит важный вклад в понимание исследуемой проблемы.\n\n"
                    "Теоретический анализ показывает многоаспектность изучаемой проблемы и необходимость "
                    "комплексного подхода к её исследованию."
                ),
                "Практическая часть": (
                    f"Практическая часть работы посвящена анализу конкретных примеров и случаев "
                    f"в контексте темы «{topic}».\n\n"
                    "На основе изученного теоретического материала был проведён анализ практических аспектов "
                    "исследуемой проблемы. Полученные данные позволяют сформулировать ряд важных выводов.\n\n"
                    "Результаты практического исследования подтверждают теоретические положения, "
                    "изложенные в предыдущей главе, и открывают новые перспективы для дальнейшего изучения темы."
                ),
                "Заключение": (
                    f"В ходе написания курсовой работы по теме «{topic}» были достигнуты поставленные цели и задачи. "
                    f"Изучены теоретические аспекты темы и проведён практический анализ.\n\n"
                    "Проведённое исследование позволило сформулировать следующие основные выводы: "
                    "тема является актуальной и требует дальнейшего изучения; "
                    f"в области {subject.lower()} существуют как устоявшиеся подходы, так и дискуссионные вопросы.\n\n"
                    "Практическая значимость работы состоит в систематизации знаний по данной теме "
                    "и формулировании рекомендаций для практического применения."
                ),
            }
            text = fallback_texts.get(section_name, f"Содержание раздела «{section_name}».")
            for paragraph in text.split("\n\n"):
                paragraph = paragraph.strip()
                if paragraph:
                    add_paragraph(doc, paragraph)

        doc.add_paragraph()

    # --- Сохранение в байты ---
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    docx_bytes = buffer.read()
    encoded = base64.b64encode(docx_bytes).decode("utf-8")

    safe_topic = topic[:40].replace(" ", "_").replace("/", "-")
    filename = f"Курсовая_{safe_topic}.docx"

    return {
        "statusCode": 200,
        "headers": {
            "Access-Control-Allow-Origin": "*",
            "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "Content-Disposition": f'attachment; filename="{filename}"',
        },
        "body": encoded,
        "isBase64Encoded": True,
    }